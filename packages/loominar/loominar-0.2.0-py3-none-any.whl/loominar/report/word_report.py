from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
from .base_report import BaseReport, SEVERITY_MAP, SEVERITY_COLORS

class WordReport(BaseReport):
    def generate(self, metrics, quality_gate, issues):
        status = quality_gate.get("status", "Report")
        doc = Document()

        # Header
        doc.add_heading("SonarQube Detailed Report", level=1)
        doc.add_paragraph(f"Project: {self.project_key}")
        doc.add_paragraph(f"Quality Gate Status: {status}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # Summary
        summary = self._build_summary(issues)
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(f"Total Issues: {summary['total']}")
        for k, v in summary["by_severity"].items():
            doc.add_paragraph(f"  {k}: {v}")
        for k, v in summary["by_type"].items():
            doc.add_paragraph(f"  {k}: {v}")

        # Charts
        pie_path, bar_path = self.generate_charts(summary)
        if pie_path and bar_path:
            doc.add_heading("Visual Summary", level=2)
            doc.add_picture(pie_path, width=Inches(4.5))
            doc.add_picture(bar_path, width=Inches(5.5))

        # Table
        doc.add_heading("Issue Details", level=2)
        if not issues:
            doc.add_paragraph("No open or confirmed issues found.")
        else:
            cols = ["Severity", "Type", "Message", "File", "Line"]
            table = doc.add_table(rows=1, cols=len(cols))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for i, col in enumerate(cols):
                hdr[i].text = col
                hdr[i].paragraphs[0].runs[0].bold = True

            for i in issues:
                row = table.add_row().cells
                sev = SEVERITY_MAP.get(i.get("severity", ""), i.get("severity", ""))
                color = SEVERITY_COLORS.get(sev, "FFFFFF")

                row[0].text = sev
                row[1].text = i.get("type", "")
                row[2].text = i.get("message", "")
                row[3].text = i.get("component", "").split(":")[-1]
                row[4].text = str(i.get("line", ""))

                tcPr = row[0]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), color)
                tcPr.append(shd)

        # Font
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        doc.add_paragraph()
        footer_text = self._render_footer()
        footer_para = doc.add_paragraph(footer_text)
        footer_para.style = doc.styles["Normal"]
        footer_para.alignment = 1

        path = self._build_filename(status)
        doc.save(path)
        print(f"✅ Word report saved: {path}")



# loominar/report/word_report.py
# Handles Word export with color-coded severities and summary