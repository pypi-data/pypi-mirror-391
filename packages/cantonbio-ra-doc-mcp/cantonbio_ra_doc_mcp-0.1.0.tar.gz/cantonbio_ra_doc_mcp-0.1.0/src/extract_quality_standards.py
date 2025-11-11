"""
Extract quality standards table from SPE-原液质量标准.docx
Specifically looking for section 4.3 检验项目、方法和标准
"""

from docx import Document
from docx.oxml.ns import qn
import re
import io
import logging
import traceback

# Get logger for this module
logger = logging.getLogger(__name__)

def extract_text_with_formatting(cell):
    """
    Extract text from cell while preserving superscript and subscript formatting with error handling
    """
    result = ""

    try:
        logger.debug(f"Processing cell with {len(cell.paragraphs)} paragraphs")

        for paragraph_idx, paragraph in enumerate(cell.paragraphs):
            try:
                logger.debug(f"Processing paragraph {paragraph_idx} with {len(paragraph.runs)} runs")

                for run_idx, run in enumerate(paragraph.runs):
                    try:
                        text = run.text
                        if text:
                            logger.debug(f"Processing run {run_idx}: '{text[:50]}...' (length: {len(text)})")

                            # Check for superscript
                            if run.font.superscript:
                                logger.debug(f"Detected superscript in run {run_idx}")
                                # Convert to superscript Unicode characters when possible
                                superscript_map = {
                                    '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴', '5': '⁵',
                                    '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹', '+': '⁺', '-': '⁻',
                                    '=': '⁼', '(': '⁽', ')': '⁾', 'n': 'ⁿ'
                                }
                                converted_text = ""
                                for char in text:
                                    converted_text += superscript_map.get(char, f"^{char}")
                                result += converted_text
                            # Check for subscript
                            elif run.font.subscript:
                                logger.debug(f"Detected subscript in run {run_idx}")
                                # Convert to subscript Unicode characters when possible
                                subscript_map = {
                                    '0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄', '5': '₅',
                                    '6': '₆', '7': '₇', '8': '₈', '9': '₉', '+': '₊', '-': '₋',
                                    '=': '₌', '(': '₍', ')': '₎', 'a': 'ₐ', 'e': 'ₑ', 'i': 'ᵢ',
                                    'o': 'ₒ', 'u': 'ᵤ', 'x': 'ₓ', 'h': 'ₕ', 'k': 'ₖ', 'l': 'ₗ',
                                    'm': 'ₘ', 'n': 'ₙ', 'p': 'ₚ', 's': 'ₛ', 't': 'ₜ'
                                }
                                converted_text = ""
                                for char in text:
                                    converted_text += subscript_map.get(char, f"_{char}")
                                result += converted_text
                            else:
                                result += text
                    except Exception as e:
                        logger.error(f"Error processing run {run_idx} in paragraph {paragraph_idx}: {str(e)}")
                        # Add the text as-is if formatting extraction fails
                        try:
                            if run.text:
                                result += run.text
                        except:
                            logger.error(f"Failed to get text from run {run_idx}")
                            continue

            except Exception as e:
                logger.error(f"Error processing paragraph {paragraph_idx}: {str(e)}")
                # Try to get paragraph text as fallback
                try:
                    if paragraph.text:
                        result += paragraph.text
                except:
                    logger.error(f"Failed to get text from paragraph {paragraph_idx}")
                    continue

            # Add space between paragraphs if there are multiple
            if len(cell.paragraphs) > 1:
                result += " "

    except Exception as e:
        logger.error(f"Critical error in extract_text_with_formatting: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        # Try to get cell text as last resort
        try:
            result = cell.text if hasattr(cell, 'text') else ""
        except:
            result = ""

    final_result = result.strip()
    logger.debug(f"Cell extraction result: '{final_result[:100]}...' (length: {len(final_result)})")
    return final_result

def extract_quality_standards_table(docx_path):
    """
    Extract quality standards table from Word document section 4.3 with detailed logging
    """
    logger.info(f"Loading document: {docx_path}")

    try:
        # Load the document
        doc = Document(docx_path)
        logger.debug(f"Document loaded successfully, contains {len(doc.paragraphs)} paragraphs")

        # Find section 4.3
        found_section_43 = False
        table_data = []

        logger.info("Searching for section 4.3 in document content...")

        # First, let's examine all paragraphs to find section 4.3
        for i, paragraph in enumerate(doc.paragraphs):
            try:
                text = paragraph.text.strip()
                if text:
                    logger.debug(f"Paragraph {i}: {text[:100]}...")

                    # Look for section 4.3
                    if re.search(r'4\.3.*检验项目.*方法.*标准', text, re.IGNORECASE):
                        logger.info(f"Found section 4.3 at paragraph {i}: {text}")
                        found_section_43 = True
                        break
            except Exception as e:
                logger.error(f"Error processing paragraph {i}: {str(e)}")
                continue

        logger.debug(f"Section 4.3 search completed. Found: {found_section_43}")

        # Now look for tables near section 4.3
        logger.info(f"Found {len(doc.tables)} tables in the document")

        for table_idx, table in enumerate(doc.tables):
            try:
                logger.debug(f"Processing table {table_idx + 1}")
                logger.debug(f"Table dimensions: {len(table.rows)} rows, {len(table.columns)} columns")

                # Extract table data with formatting
                table_content = []
                for row_idx, row in enumerate(table.rows):
                    try:
                        row_data = []
                        for cell_idx, cell in enumerate(row.cells):
                            try:
                                # Use new function to extract text with formatting
                                cell_text = extract_text_with_formatting(cell)
                                # Clean up excessive whitespace but preserve formatting
                                cell_text = ' '.join(cell_text.split())
                                row_data.append(cell_text)
                            except Exception as e:
                                logger.error(f"Error processing cell [{row_idx}, {cell_idx}]: {str(e)}")
                                row_data.append("")  # Add empty string for failed cells

                        table_content.append(row_data)
                        logger.debug(f"Row {row_idx}: {row_data}")
                    except Exception as e:
                        logger.error(f"Error processing row {row_idx} in table {table_idx + 1}: {str(e)}")
                        continue

                # Check if this table contains quality standards data
                if table_content:
                    header_row = table_content[0] if table_content else []
                    header_text = ' '.join(header_row).lower()

                    # Look for keywords that indicate this is the quality standards table
                    quality_keywords = ['检验项目', '检验方法', '质量标准', '类型', '项目', '方法', '标准']
                    keyword_count = sum(1 for keyword in quality_keywords if keyword in header_text)

                    logger.debug(f"Table {table_idx + 1} header keywords found: {keyword_count}")
                    if keyword_count >= 2:  # At least 2 keywords match
                        logger.info(f"Table {table_idx + 1} appears to be the quality standards table (keyword matches: {keyword_count})")
                        table_data = table_content
                        break
                else:
                    logger.warning(f"Table {table_idx + 1} has no content")

            except Exception as e:
                logger.error(f"Error processing table {table_idx + 1}: {str(e)}")
                logger.error(f"Table processing traceback: {traceback.format_exc()}")
                continue

        if not table_data:
            logger.warning("No quality standards table found in document")

        return table_data

    except Exception as e:
        logger.error(f"Critical error in extract_quality_standards_table: {str(e)}")
        logger.error(f"Full traceback: {traceback.format_exc()}")
        raise

def format_as_markdown_table(table_data, target_columns=['类型', '检验项目', '检验方法', '质量标准']):
    """
    Convert table data to markdown format
    """
    if not table_data:
        return "No table data found"

    # Try to map columns to target format
    header_row = table_data[0] if table_data else []
    print(f"\nOriginal headers: {header_row}")

    # Create markdown table
    markdown_lines = []

    # Header
    markdown_lines.append(f"| {' | '.join(target_columns)} |")
    markdown_lines.append(f"| {' | '.join(['---'] * len(target_columns))} |")

    # Data rows
    for row in table_data[1:]:  # Skip header row
        # Pad row to match target columns length
        padded_row = row + [''] * (len(target_columns) - len(row))
        # Truncate if too long
        padded_row = padded_row[:len(target_columns)]

        markdown_lines.append(f"| {' | '.join(padded_row)} |")

    return '\n'.join(markdown_lines)

# def extract_quality_standards_table_from_bytes(file_bytes: bytes) -> str:
#     """
#     Extract quality standards table from Word document bytes

#     Args:
#         file_bytes: Binary content of the Word document

#     Returns:
#         Markdown formatted table string
#     """
#     print("Extracting quality standards table from Word document bytes...")

#     try:
#         doc = Document(io.BytesIO(file_bytes))

#         found_section_43 = False
#         table_data = []

#         print("Searching through document content...")

#         for i, paragraph in enumerate(doc.paragraphs):
#             text = paragraph.text.strip()
#             if text:
#                 print(f"Paragraph {i}: {text[:100]}...")

#                 if re.search(r'4\.3.*检验项目.*方法.*标准', text, re.IGNORECASE):
#                     print(f"Found section 4.3 at paragraph {i}: {text}")
#                     found_section_43 = True
#                     break

#         print(f"\nFound {len(doc.tables)} tables in the document")

#         for table_idx, table in enumerate(doc.tables):
#             print(f"\n--- Table {table_idx + 1} ---")
#             print(f"Rows: {len(table.rows)}, Columns: {len(table.columns)}")

#             table_content = []
#             for row_idx, row in enumerate(table.rows):
#                 row_data = []
#                 for cell in row.cells:
#                     cell_text = extract_text_with_formatting(cell)
#                     cell_text = ' '.join(cell_text.split())
#                     row_data.append(cell_text)
#                 table_content.append(row_data)
#                 print(f"Row {row_idx}: {row_data}")

#             if table_content:
#                 header_row = table_content[0] if table_content else []
#                 header_text = ' '.join(header_row).lower()

#                 quality_keywords = ['检验项目', '检验方法', '质量标准', '类型', '项目', '方法', '标准']
#                 keyword_count = sum(1 for keyword in quality_keywords if keyword in header_text)

#                 if keyword_count >= 2:
#                     print(f"Table {table_idx + 1} appears to be the quality standards table (keyword matches: {keyword_count})")
#                     table_data = table_content
#                     break

#         if table_data:
#             print(f"\n=== Extracted Table Data ({len(table_data)} rows) ===")
#             for i, row in enumerate(table_data):
#                 print(f"Row {i}: {row}")

#             print(f"\n=== Markdown Format ===")
#             markdown_table = format_as_markdown_table(table_data)
#             return markdown_table
#         else:
#             return "No quality standards table found in the document"

#     except Exception as e:
#         error_msg = f"Error processing document: {e}"
#         print(error_msg)
#         import traceback
#         traceback.print_exc()
#         return error_msg

def extract_quality_standards_table_from_docx(docx_path: str):
    """
    Extract quality standards table from Word document with comprehensive error handling and debugging
    """
    logger.info(f"Starting extraction from document: {docx_path}")

    try:
        # Check if file exists
        import os
        if not os.path.exists(docx_path):
            error_msg = f"File not found: {docx_path}"
            logger.error(error_msg)
            return f"Error: {error_msg}"

        logger.debug(f"File exists, size: {os.path.getsize(docx_path)} bytes")

        # Extract table data
        table_data = extract_quality_standards_table(docx_path)

        if table_data:
            logger.info(f"Successfully extracted table with {len(table_data)} rows")
            logger.debug("Extracted table content:")
            for i, row in enumerate(table_data):
                logger.debug(f"Row {i}: {row}")

            # Convert to markdown
            markdown_table = format_as_markdown_table(table_data)
            logger.debug(f"Generated markdown table:\n{markdown_table}")
            return markdown_table
        else:
            error_msg = "No quality standards table found in the document"
            logger.warning(error_msg)
            return f"Warning: {error_msg}"

    except Exception as e:
        error_msg = f"Error processing document {docx_path}: {str(e)}"
        logger.error(error_msg)
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return f"Error: {error_msg}"
