"""
Unit tests for fill_quality_standards module
"""

import pytest
import tempfile
import os
from unittest.mock import patch, MagicMock
from docx import Document
from src.fill_quality_standards import (
    parse_markdown_table_from_string,
    parse_markdown_table_from_file,
    restore_formatting_to_cell,
    clear_table_content,
    insert_table_rows,
    find_quality_standards_table,
    fill_word_document_table,
    fill_quality_standards_from_markdown,
    fill_quality_standards_from_file,
    fill_quality_standards_inplace,
    merge_cells_in_column
)


class TestMarkdownParsing:
    """Test markdown table parsing functions"""

    def test_parse_markdown_table_from_string_valid(self):
        """Test parsing valid markdown table"""
        markdown_content = """
| 类型 | 检验项目 | 检验方法 | 质量标准 |
|------|----------|----------|----------|
| 理化检定 | 颜色 | 目视法 | 无色澄清溶液 |
| 理化检定 | pH值 | pH计法 | 6.5-8.5 |
| 鉴别 | 蛋白质鉴别 | SDS-PAGE | 符合规定 |
"""
        result = parse_markdown_table_from_string(markdown_content)

        assert len(result) == 3
        assert result[0] == ['理化检定', '颜色', '目视法', '无色澄清溶液']
        assert result[1] == ['理化检定', 'pH值', 'pH计法', '6.5-8.5']
        assert result[2] == ['鉴别', '蛋白质鉴别', 'SDS-PAGE', '符合规定']

    def test_parse_markdown_table_from_string_empty(self):
        """Test parsing empty markdown content"""
        result = parse_markdown_table_from_string("")
        assert result == []

    def test_parse_markdown_table_from_string_no_data_rows(self):
        """Test parsing markdown with only header"""
        markdown_content = """
| 类型 | 检验项目 | 检验方法 | 质量标准 |
|------|----------|----------|----------|
"""
        result = parse_markdown_table_from_string(markdown_content)
        assert result == []

    def test_parse_markdown_table_from_string_with_unicode_formatting(self):
        """Test parsing markdown with Unicode super/subscript"""
        markdown_content = """
| 类型 | 检验项目 | 检验方法 | 质量标准 |
|------|----------|----------|----------|
| 含量 | Ca²⁺含量 | ICP-MS | ≥95% |
| 纯度 | H₂O含量 | Karl Fischer | ≤5% |
"""
        result = parse_markdown_table_from_string(markdown_content)

        assert len(result) == 2
        assert result[0] == ['含量', 'Ca²⁺含量', 'ICP-MS', '≥95%']
        assert result[1] == ['纯度', 'H₂O含量', 'Karl Fischer', '≤5%']

    def test_parse_markdown_table_from_file(self):
        """Test parsing markdown table from file"""
        markdown_content = """
| 类型 | 检验项目 | 检验方法 | 质量标准 |
|------|----------|----------|----------|
| 理化检定 | 颜色 | 目视法 | 无色澄清溶液 |
"""

        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
            f.write(markdown_content)
            temp_file = f.name

        try:
            result = parse_markdown_table_from_file(temp_file)
            assert len(result) == 1
            assert result[0] == ['理化检定', '颜色', '目视法', '无色澄清溶液']
        finally:
            os.unlink(temp_file)

    def test_parse_markdown_table_from_file_not_found(self):
        """Test parsing markdown table from non-existent file"""
        with pytest.raises(Exception):
            parse_markdown_table_from_file("non_existent_file.md")


class TestWordDocumentManipulation:
    """Test Word document manipulation functions"""

    def test_clear_table_content_keep_header(self):
        """Test clearing table content while keeping header"""
        # Create a mock document with a table
        doc = Document()
        table = doc.add_table(rows=3, cols=2)

        # Add some content
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"
        table.cell(2, 0).text = "Data 3"
        table.cell(2, 1).text = "Data 4"

        # Clear content keeping header
        clear_table_content(table, keep_header=True)

        # Should have only header row
        assert len(table.rows) == 1
        assert table.cell(0, 0).text == "Header 1"
        assert table.cell(0, 1).text == "Header 2"

    def test_clear_table_content_remove_all(self):
        """Test clearing all table content"""
        doc = Document()
        table = doc.add_table(rows=3, cols=2)

        clear_table_content(table, keep_header=False)

        # Should have no rows
        assert len(table.rows) == 0

    def test_insert_table_rows(self):
        """Test inserting rows into table"""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)  # Start with header row

        insert_table_rows(table, 3)

        # Should have 4 rows total (1 header + 3 new)
        assert len(table.rows) == 4

    def test_find_quality_standards_table_found(self):
        """Test finding quality standards table"""
        doc = Document()

        # Add a regular table
        table1 = doc.add_table(rows=1, cols=2)
        table1.cell(0, 0).text = "Name"
        table1.cell(0, 1).text = "Value"

        # Add quality standards table
        table2 = doc.add_table(rows=1, cols=4)
        table2.cell(0, 0).text = "类型"
        table2.cell(0, 1).text = "检验项目"
        table2.cell(0, 2).text = "检验方法"
        table2.cell(0, 3).text = "质量标准"

        result = find_quality_standards_table(doc)
        assert result == 1  # Second table (index 1)

    def test_find_quality_standards_table_not_found(self):
        """Test when quality standards table is not found"""
        doc = Document()

        # Add a regular table
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"

        result = find_quality_standards_table(doc)
        assert result is None

    @patch('src.fill_quality_standards.logger')
    def test_restore_formatting_to_cell_with_unicode(self, mock_logger):
        """Test restoring formatting with Unicode characters"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        # Test with superscript Unicode
        restore_formatting_to_cell(cell, "Ca²⁺")

        # Should have content (exact formatting test is complex due to docx internals)
        assert len(cell.paragraphs) > 0
        assert len(cell.paragraphs[0].runs) > 0

    @patch('src.fill_quality_standards.logger')
    def test_restore_formatting_to_cell_with_notation(self, mock_logger):
        """Test restoring formatting with ^{} and _{} notation"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        # Test with ^{} notation
        restore_formatting_to_cell(cell, "H^{+}")

        assert len(cell.paragraphs) > 0
        assert len(cell.paragraphs[0].runs) > 0

    @patch('src.fill_quality_standards.logger')
    def test_restore_formatting_to_cell_plain_text(self, mock_logger):
        """Test restoring formatting with plain text"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        restore_formatting_to_cell(cell, "Plain text")

        assert len(cell.paragraphs) > 0
        # Check that the text was processed (there should be runs)
        assert len(cell.paragraphs[0].runs) > 0
        # Combine all run texts to check full content
        full_text = ''.join(run.text for run in cell.paragraphs[0].runs)
        assert full_text == "Plain text"

    def test_merge_cells_in_column_no_duplicates(self):
        """Test that merging cells doesn't create duplicate content"""
        doc = Document()
        table = doc.add_table(rows=4, cols=2)

        # Fill cells with identical content
        for i in range(1, 4):  # Skip header row
            table.cell(i, 0).text = "重复内容"
            table.cell(i, 1).text = f"不同内容{i}"

        # Test merging first column (rows 1-3)
        merge_cells_in_column(table, 0, 1, 3)

        # Verify no duplicate content
        merged_cell_text = table.cell(1, 0).text.strip()
        assert merged_cell_text == "重复内容", f"Expected '重复内容', got '{merged_cell_text}'"
        assert '\n' not in merged_cell_text, f"Found duplicate content: '{merged_cell_text}'"

        # Verify other cells are not affected
        assert table.cell(1, 1).text == "不同内容1"
        assert table.cell(2, 1).text == "不同内容2"
        assert table.cell(3, 1).text == "不同内容3"


class TestIntegrationFunctions:
    """Test integration functions"""

    def test_fill_word_document_table_success(self):
        """Test successful filling of Word document table"""
        # Create a test document
        doc = Document()
        table = doc.add_table(rows=1, cols=4)
        table.cell(0, 0).text = "类型"
        table.cell(0, 1).text = "检验项目"
        table.cell(0, 2).text = "检验方法"
        table.cell(0, 3).text = "质量标准"

        # Save to temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_doc:
            doc.save(temp_doc.name)
            input_path = temp_doc.name

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_output:
            output_path = temp_output.name

        try:
            # Test data
            table_data = [
                ['理化检定', '颜色', '目视法', '无色澄清溶液'],
                ['理化检定', 'pH值', 'pH计法', '6.5-8.5']
            ]

            result = fill_word_document_table(
                input_path, output_path, table_data, table_index=0, auto_merge=False
            )

            assert "Successfully filled table" in result
            assert os.path.exists(output_path)

            # Verify the filled document
            filled_doc = Document(output_path)
            filled_table = filled_doc.tables[0]

            assert len(filled_table.rows) == 3  # Header + 2 data rows
            assert filled_table.cell(1, 0).text == '理化检定'
            assert filled_table.cell(1, 1).text == '颜色'
            assert filled_table.cell(2, 0).text == '理化检定'
            assert filled_table.cell(2, 1).text == 'pH值'

        finally:
            # Cleanup
            if os.path.exists(input_path):
                os.unlink(input_path)
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_fill_word_document_table_no_table_found(self):
        """Test when no suitable table is found"""
        # Create a document without quality standards table
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_doc:
            doc.save(temp_doc.name)
            input_path = temp_doc.name

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_output:
            output_path = temp_output.name

        try:
            table_data = [['理化检定', '颜色', '目视法', '无色澄清溶液']]

            result = fill_word_document_table(
                input_path, output_path, table_data, table_index=None
            )

            assert "Error: No quality standards table found" in result

        finally:
            if os.path.exists(input_path):
                os.unlink(input_path)
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_fill_word_document_table_invalid_index(self):
        """Test with invalid table index"""
        doc = Document()
        table = doc.add_table(rows=1, cols=4)

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_doc:
            doc.save(temp_doc.name)
            input_path = temp_doc.name

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_output:
            output_path = temp_output.name

        try:
            table_data = [['理化检定', '颜色', '目视法', '无色澄清溶液']]

            result = fill_word_document_table(
                input_path, output_path, table_data, table_index=99
            )

            assert "Error: Table index 99 not found" in result

        finally:
            if os.path.exists(input_path):
                os.unlink(input_path)
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_fill_quality_standards_from_markdown(self):
        """Test filling from markdown content"""
        # Create test document
        doc = Document()
        table = doc.add_table(rows=1, cols=4)
        table.cell(0, 0).text = "类型"
        table.cell(0, 1).text = "检验项目"
        table.cell(0, 2).text = "检验方法"
        table.cell(0, 3).text = "质量标准"

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_doc:
            doc.save(temp_doc.name)
            input_path = temp_doc.name

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_output:
            output_path = temp_output.name

        try:
            markdown_content = """
| 类型 | 检验项目 | 检验方法 | 质量标准 |
|------|----------|----------|----------|
| 理化检定 | 颜色 | 目视法 | 无色澄清溶液 |
| 理化检定 | pH值 | pH计法 | 6.5-8.5 |
"""

            result = fill_quality_standards_from_markdown(
                input_path, output_path, markdown_content, table_index=0
            )

            assert "Successfully filled table" in result

        finally:
            if os.path.exists(input_path):
                os.unlink(input_path)
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_fill_quality_standards_from_markdown_invalid_content(self):
        """Test filling from invalid markdown content"""
        doc = Document()
        table = doc.add_table(rows=1, cols=4)

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_doc:
            doc.save(temp_doc.name)
            input_path = temp_doc.name

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_output:
            output_path = temp_output.name

        try:
            # Invalid markdown content (no data rows)
            markdown_content = """
| 类型 | 检验项目 | 检验方法 | 质量标准 |
|------|----------|----------|----------|
"""

            result = fill_quality_standards_from_markdown(
                input_path, output_path, markdown_content
            )

            assert "Error: No valid table data found" in result

        finally:
            if os.path.exists(input_path):
                os.unlink(input_path)
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_fill_quality_standards_from_file_success(self):
        """Test filling from markdown file"""
        # Create test document
        doc = Document()
        table = doc.add_table(rows=1, cols=4)
        table.cell(0, 0).text = "类型"
        table.cell(0, 1).text = "检验项目"
        table.cell(0, 2).text = "检验方法"
        table.cell(0, 3).text = "质量标准"

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_doc:
            doc.save(temp_doc.name)
            input_path = temp_doc.name

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_output:
            output_path = temp_output.name

        # Create markdown file
        markdown_content = """
| 类型 | 检验项目 | 检验方法 | 质量标准 |
|------|----------|----------|----------|
| 理化检定 | 颜色 | 目视法 | 无色澄清溶液 |
"""

        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as temp_md:
            temp_md.write(markdown_content)
            md_path = temp_md.name

        try:
            result = fill_quality_standards_from_file(
                input_path, output_path, md_path, table_index=0
            )

            assert "Successfully filled table" in result

        finally:
            for path in [input_path, output_path, md_path]:
                if os.path.exists(path):
                    os.unlink(path)


class TestErrorHandling:
    """Test error handling scenarios"""

    def test_fill_word_document_table_file_not_found(self):
        """Test handling of non-existent input file"""
        result = fill_word_document_table(
            "non_existent.docx", "output.docx", [['test', 'data']]
        )

        assert "Error filling Word document table:" in result

    def test_fill_word_document_table_empty_data(self):
        """Test handling of empty table data"""
        doc = Document()
        table = doc.add_table(rows=1, cols=4)
        # Add headers to make it a quality standards table
        table.cell(0, 0).text = "类型"
        table.cell(0, 1).text = "检验项目"
        table.cell(0, 2).text = "检验方法"
        table.cell(0, 3).text = "质量标准"

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_doc:
            doc.save(temp_doc.name)
            input_path = temp_doc.name

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_output:
            output_path = temp_output.name

        try:
            result = fill_word_document_table(input_path, output_path, [])

            assert "Error: No table data provided" in result

        finally:
            if os.path.exists(input_path):
                os.unlink(input_path)
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_fill_quality_standards_from_file_file_not_found(self):
        """Test handling of non-existent markdown file"""
        result = fill_quality_standards_from_file(
            "input.docx", "output.docx", "non_existent.md"
        )

        assert "Error processing markdown file:" in result

    def test_fill_quality_standards_inplace_success(self):
        """Test filling document in-place"""
        # Create test document
        doc = Document()
        table = doc.add_table(rows=1, cols=4)
        table.cell(0, 0).text = "类型"
        table.cell(0, 1).text = "检验项目"
        table.cell(0, 2).text = "检验方法"
        table.cell(0, 3).text = "质量标准"

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_doc:
            doc.save(temp_doc.name)
            doc_path = temp_doc.name

        try:
            markdown_content = """
| 类型 | 检验项目 | 检验方法 | 质量标准 |
|------|----------|----------|----------|
| 理化检定 | 颜色 | 目视法 | 无色澄清溶液 |
| 理化检定 | pH值 | pH计法 | 6.5-8.5 |
"""

            result = fill_quality_standards_inplace(
                doc_path, markdown_content, table_index=0
            )

            assert "Successfully filled table" in result

            # Verify the document was modified in-place
            modified_doc = Document(doc_path)
            modified_table = modified_doc.tables[0]

            assert len(modified_table.rows) == 3  # Header + 2 data rows

            # Verify cell content after merging (should not have duplicates)
            cell_1_0_text = modified_table.cell(1, 0).text.strip()
            cell_2_0_text = modified_table.cell(2, 0).text.strip()

            # The merged cell should contain the content only once, not duplicated
            assert cell_1_0_text == '理化检定', f"Expected '理化检定', got '{cell_1_0_text}'"
            assert modified_table.cell(1, 1).text == '颜色'

            # Due to merging, the second row's first cell might be empty or same as first
            # but it should not contain duplicate text
            assert '理化检定' in cell_2_0_text or cell_2_0_text == '', f"Unexpected content in merged cell: '{cell_2_0_text}'"
            assert modified_table.cell(2, 1).text == 'pH值'

            # Most importantly: verify no duplicate content in merged cells
            assert '\n理化检定' not in cell_1_0_text, f"Found duplicate content in merged cell: '{cell_1_0_text}'"

        finally:
            if os.path.exists(doc_path):
                os.unlink(doc_path)


if __name__ == "__main__":
    pytest.main([__file__])