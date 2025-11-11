#!/usr/bin/env python3
"""
Example usage of fill_quality_standards module

This example demonstrates how to:
1. Parse markdown table data
2. Fill a Word document table with the data
3. Handle Unicode formatting (superscript/subscript)
4. Automatically merge duplicate cells
"""

import tempfile
import os
import sys
sys.path.append('..')
from docx import Document
from src.fill_quality_standards import (
    fill_quality_standards_from_markdown,
    fill_quality_standards_inplace,
    parse_markdown_table_from_string
)

def create_example_template():
    """Create an example Word document with quality standards table template"""
    doc = Document()

    # Add title
    title = doc.add_heading('质量标准', 1)

    # Add some introductory text
    doc.add_paragraph('本文档包含药品的质量标准检验项目、方法和标准。')

    # Create table with headers
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Set headers
    headers = ['类型', '检验项目', '检验方法', '质量标准']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Make headers bold
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    return doc

def main():
    """Main example function"""
    print("=== 质量标准表格填充示例 ===\n")

    # Step 1: Create example markdown table data
    markdown_content = """
| 类型 | 检验项目 | 检验方法 | 质量标准 |
|------|----------|----------|----------|
| 理化检定 | 颜色 | 目视法 | 无色澄清溶液 |
| 理化检定 | 澄清度 | 目视法 | 澄清 |
| 理化检定 | pH值 | pH计法 | 6.5-8.5 |
| 鉴别 | 蛋白质鉴别 | SDS-PAGE | 符合规定 |
| 鉴别 | 免疫印迹法 | Western Blot | 符合规定 |
| 含量 | 蛋白质含量 | Bradford法 | ≥95% |
| 含量 | Ca²⁺含量 | ICP-MS | 符合规定 |
| 纯度和杂质 | 相关蛋白质 | SEC-HPLC | ≤5% |
| 纯度和杂质 | 宿主细胞蛋白 | ELISA | ≤100ng/mg |
| 纯度和杂质 | 宿主细胞DNA | qPCR | ≤10ng/mg |
| 纯度和杂质 | 内毒素 | LAL法 | ≤0.25EU/mg |
| 安全性 | 无菌检查 | 薄膜过滤法 | 符合规定 |
| 安全性 | 支原体 | PCR法 | 阴性 |
"""

    print("1. 解析Markdown表格数据...")
    table_data = parse_markdown_table_from_string(markdown_content)
    print(f"   成功解析 {len(table_data)} 行数据")

    # Show sample data
    print("   示例数据:")
    for i, row in enumerate(table_data[:3]):  # Show first 3 rows
        print(f"     Row {i+1}: {row}")
    if len(table_data) > 3:
        print(f"     ... 还有 {len(table_data) - 3} 行")
    print()

    # Step 2: Create template document
    print("2. 创建Word文档模板...")
    template_doc = create_example_template()

    # Save template to temporary file
    with tempfile.NamedTemporaryFile(suffix='_template.docx', delete=False) as temp_template:
        template_doc.save(temp_template.name)
        template_path = temp_template.name

    print(f"   模板文档已创建: {template_path}")

    # Step 3: Fill the template with data (in-place modification)
    print("3. 填充质量标准表格 (就地修改)...")

    try:
        result = fill_quality_standards_inplace(
            template_path,
            markdown_content,
            table_index=0,  # Use the first (and only) table
            auto_merge=True  # Automatically merge duplicate cells
        )

        print(f"   填充结果: {result}")

        # Step 4: Verify the result
        print("4. 验证填充结果...")
        filled_doc = Document(template_path)  # Document was modified in-place
        filled_table = filled_doc.tables[0]

        print(f"   表格总行数: {len(filled_table.rows)} (包含标题行)")
        print(f"   表格总列数: {len(filled_table.columns)}")

        # Show first few filled rows
        print("   填充的数据 (前5行):")
        for i in range(1, min(6, len(filled_table.rows))):  # Skip header
            row = filled_table.rows[i]
            row_data = [cell.text.strip() for cell in row.cells]
            print(f"     Row {i}: {' | '.join(row_data)}")

        print(f"\n✅ 成功就地修改文档: {template_path}")
        print(f"✅ 可以打开该文档查看填充和合并效果")

    except Exception as e:
        print(f"❌ 填充过程中发生错误: {e}")
        import traceback
        traceback.print_exc()

    finally:
        # Note: In a real application, you might want to clean up temp files
        # For this example, we keep them so you can inspect the results
        print(f"\n📁 修改后的文档:")
        print(f"   文档路径: {template_path}")
        print(f"   (可以手动删除这个临时文件)")

if __name__ == "__main__":
    main()