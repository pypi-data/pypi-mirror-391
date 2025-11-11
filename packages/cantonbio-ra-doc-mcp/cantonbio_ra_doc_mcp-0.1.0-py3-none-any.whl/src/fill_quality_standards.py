"""
Fill Word document table with quality standards data from markdown format
基于fill_word_table.py和merge_word_cells.py的功能，实现将markdown格式的质量标准表格填写到指定docx文件中
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import logging
import traceback
from typing import List, Tuple, Optional, Dict

# Get logger for this module
logger = logging.getLogger(__name__)

def parse_markdown_table_from_string(markdown_content: str) -> List[List[str]]:
    """
    Parse markdown table from string content to extract data

    Args:
        markdown_content: Markdown table content as string

    Returns:
        List of rows, each row is a list of cell values
    """
    logger.debug(f"Parsing markdown table from string content")

    data = []
    lines = markdown_content.strip().split('\n')

    # Find table lines (skip header separator line with ---)
    table_started = False
    header_found = False

    for line_num, line in enumerate(lines):
        line = line.strip()
        logger.debug(f"Processing line {line_num}: '{line}'")

        if line.startswith('|') and line.endswith('|'):
            if '---' in line:
                # This is the separator line, skip it
                logger.debug(f"Skipping separator line: {line}")
                continue

            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove first and last empty elements

            if not header_found:
                # This is the header row
                logger.debug(f"Found header row: {cells}")
                header_found = True
                # Store header for reference but don't include in data
                continue
            else:
                # This is a data row
                if cells and any(cell.strip() for cell in cells):  # Skip empty rows
                    logger.debug(f"Found data row: {cells}")
                    data.append(cells)

    logger.info(f"Parsed {len(data)} data rows from markdown table")
    return data

def parse_markdown_table_from_file(md_file_path: str) -> List[List[str]]:
    """
    Parse markdown table from file to extract data

    Args:
        md_file_path: Path to markdown file

    Returns:
        List of rows, each row is a list of cell values
    """
    logger.info(f"Reading markdown table from file: {md_file_path}")

    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return parse_markdown_table_from_string(content)
    except Exception as e:
        logger.error(f"Error reading markdown file {md_file_path}: {str(e)}")
        raise

def restore_formatting_to_cell(cell, text: str):
    """
    Restore superscript and subscript formatting from Unicode characters to Word formatting

    Args:
        cell: Word table cell
        text: Text with Unicode super/subscript characters
    """
    try:
        # Clear existing content
        cell.text = ""
        paragraph = cell.paragraphs[0]

        # Maps for converting Unicode back to normal characters
        superscript_map = {
            '⁰': '0', '¹': '1', '²': '2', '³': '3', '⁴': '4', '⁵': '5',
            '⁶': '6', '⁷': '7', '⁸': '8', '⁹': '9', '⁺': '+', '⁻': '-',
            '⁼': '=', '⁽': '(', '⁾': ')', 'ⁿ': 'n'
        }

        subscript_map = {
            '₀': '0', '₁': '1', '₂': '2', '₃': '3', '₄': '4', '₅': '5',
            '₆': '6', '₇': '7', '₈': '8', '₉': '9', '₊': '+', '₋': '-',
            '₌': '=', '₍': '(', '₎': ')', 'ₐ': 'a', 'ₑ': 'e', 'ᵢ': 'i',
            'ₒ': 'o', 'ᵤ': 'u', 'ₓ': 'x', 'ₕ': 'h', 'ₖ': 'k', 'ₗ': 'l',
            'ₘ': 'm', 'ₙ': 'n', 'ₚ': 'p', 'ₛ': 's', 'ₜ': 't'
        }

        # Also handle ^{} and _{} notation
        # Pattern for ^{text} and _{text}
        super_pattern = r'\^\{([^}]+)\}'
        sub_pattern = r'_\{([^}]+)\}'

        # Replace ^{text} and _{text} patterns first
        def replace_super(match):
            return ''.join(superscript_map.get(c, f'^{c}') for c in match.group(1))

        def replace_sub(match):
            return ''.join(subscript_map.get(c, f'_{c}') for c in match.group(1))

        text = re.sub(super_pattern, replace_super, text)
        text = re.sub(sub_pattern, replace_sub, text)

        # Process character by character to handle formatting
        i = 0
        while i < len(text):
            char = text[i]

            if char in superscript_map:
                # Add superscript character
                run = paragraph.add_run(superscript_map[char])
                run.font.superscript = True
            elif char in subscript_map:
                # Add subscript character
                run = paragraph.add_run(subscript_map[char])
                run.font.subscript = True
            elif char == '^' and i + 1 < len(text):
                # Handle ^character notation
                i += 1
                next_char = text[i]
                run = paragraph.add_run(next_char)
                run.font.superscript = True
            elif char == '_' and i + 1 < len(text):
                # Handle _character notation
                i += 1
                next_char = text[i]
                run = paragraph.add_run(next_char)
                run.font.subscript = True
            else:
                # Normal character
                run = paragraph.add_run(char)

            i += 1

    except Exception as e:
        logger.warning(f"Error restoring formatting to cell, using plain text: {str(e)}")
        # Fallback to plain text
        cell.text = text

def clear_table_content(table, keep_header: bool = True):
    """
    Clear all rows in the table except optionally the header row

    Args:
        table: Word table object
        keep_header: Whether to keep the first row (header)
    """
    start_row = 1 if keep_header else 0
    rows_to_remove = len(table.rows) - start_row

    logger.debug(f"Clearing {rows_to_remove} existing rows (keeping header: {keep_header})...")

    # Remove rows from the end to avoid index issues
    for _ in range(rows_to_remove):
        if len(table.rows) > start_row:
            # Remove the last row
            table._tbl.remove(table.rows[-1]._tr)

    logger.debug(f"Table now has {len(table.rows)} rows")

def insert_table_rows(table, num_rows: int):
    """
    Insert the specified number of empty rows into the table

    Args:
        table: Word table object
        num_rows: Number of rows to insert
    """
    logger.debug(f"Inserting {num_rows} new empty rows...")

    for _ in range(num_rows):
        # Add a new row to the table
        new_row = table.add_row()
        # Initialize cells with empty strings
        for cell in new_row.cells:
            cell.text = ""

    logger.debug(f"Table now has {len(table.rows)} total rows")

def find_quality_standards_table(doc: Document) -> Optional[int]:
    """
    Find the quality standards table in the document

    Args:
        doc: Word document object

    Returns:
        Table index if found, None otherwise
    """
    logger.debug(f"Searching for quality standards table in {len(doc.tables)} tables")

    for table_idx, table in enumerate(doc.tables):
        try:
            if len(table.rows) > 0:
                # Check header row for quality standards keywords
                header_row = table.rows[0]
                header_text = ' '.join(cell.text for cell in header_row.cells).lower()

                quality_keywords = ['检验项目', '检验方法', '质量标准', '类型', '项目', '方法', '标准']
                keyword_count = sum(1 for keyword in quality_keywords if keyword in header_text)

                logger.debug(f"Table {table_idx}: header keywords found: {keyword_count}")

                if keyword_count >= 2:  # At least 2 keywords match
                    logger.info(f"Found quality standards table at index {table_idx}")
                    return table_idx

        except Exception as e:
            logger.warning(f"Error checking table {table_idx}: {str(e)}")
            continue

    logger.warning("No quality standards table found")
    return None

def merge_cells_in_column(table, col_index: int, start_row: int, end_row: int):
    """
    Merge cells in a column from start_row to end_row (inclusive)

    Note: This function prevents duplicate content by clearing subsequent cells before merging

    Args:
        table: Word table object
        col_index: Column index to merge
        start_row: Starting row index
        end_row: Ending row index
    """
    if start_row >= end_row or end_row >= len(table.rows):
        logger.debug(f"Skipping merge for column {col_index}, rows {start_row}-{end_row} (invalid range)")
        return

    try:
        logger.debug(f"Merging column {col_index}, rows {start_row}-{end_row}")

        # Get the first cell to merge into
        first_cell = table.rows[start_row].cells[col_index]

        # Store the original content from the first cell
        original_content = first_cell.text.strip()

        # Clear content from cells to be merged to prevent duplication
        for row_idx in range(start_row + 1, end_row + 1):
            if row_idx < len(table.rows):
                cell = table.rows[row_idx].cells[col_index]
                cell.text = ""  # Clear content before merging

        # Now merge the cells (they're empty, so no duplicate content)
        for row_idx in range(start_row + 1, end_row + 1):
            if row_idx < len(table.rows):
                cell_to_merge = table.rows[row_idx].cells[col_index]
                first_cell.merge(cell_to_merge)

        # Ensure the merged cell has the correct content
        if first_cell.text.strip() != original_content:
            first_cell.text = original_content

        logger.debug(f"Successfully merged column {col_index}, rows {start_row}-{end_row}")

    except Exception as e:
        logger.warning(f"Could not merge cells {start_row}-{end_row} in column {col_index}: {str(e)}")
        # If merging fails, at least clear the duplicate text in subsequent cells
        try:
            for row_idx in range(start_row + 1, end_row + 1):
                if row_idx < len(table.rows):
                    cell = table.rows[row_idx].cells[col_index]
                    cell.text = ""
        except Exception as e2:
            logger.warning(f"Could not clear duplicate text: {str(e2)}")

def auto_merge_duplicate_cells(table, target_columns: List[str] = ['类型', '检验项目']):
    """
    Automatically merge cells with duplicate content in specified columns

    Args:
        table: Word table object
        target_columns: List of column names to check for merging
    """
    if len(table.rows) <= 1:
        logger.debug("Not enough rows for merging")
        return

    # Get header row to find column indices
    header_row = table.rows[0]
    column_mapping = {}

    for col_idx, cell in enumerate(header_row.cells):
        header_text = cell.text.strip()
        for target_col in target_columns:
            if target_col in header_text:
                column_mapping[target_col] = col_idx
                break

    logger.debug(f"Column mapping for merging: {column_mapping}")

    # Merge 类型 column
    if '类型' in column_mapping:
        col_idx = column_mapping['类型']
        logger.debug(f"Processing 类型 column (index {col_idx}) for merging")

        current_type = ""
        merge_start = -1

        for row_idx in range(1, len(table.rows)):  # Skip header
            cell_text = table.rows[row_idx].cells[col_idx].text.strip()

            if cell_text != current_type:
                # Different type found, merge previous group if needed
                if merge_start != -1 and row_idx - merge_start > 1:
                    merge_cells_in_column(table, col_idx, merge_start, row_idx - 1)

                current_type = cell_text
                merge_start = row_idx

        # Handle last group
        if merge_start != -1 and len(table.rows) - merge_start > 1:
            merge_cells_in_column(table, col_idx, merge_start, len(table.rows) - 1)

    # Merge 检验项目 column within same 类型
    if '检验项目' in column_mapping and '类型' in column_mapping:
        type_col_idx = column_mapping['类型']
        item_col_idx = column_mapping['检验项目']

        logger.debug(f"Processing 检验项目 column (index {item_col_idx}) for merging within same 类型")

        current_type = ""
        current_item = ""
        item_start = -1

        for row_idx in range(1, len(table.rows)):
            type_text = table.rows[row_idx].cells[type_col_idx].text.strip()
            item_text = table.rows[row_idx].cells[item_col_idx].text.strip()

            if type_text != current_type:
                # Different type, merge previous item group if needed
                if item_start != -1 and row_idx - item_start > 1 and current_item:
                    merge_cells_in_column(table, item_col_idx, item_start, row_idx - 1)

                current_type = type_text
                current_item = item_text
                item_start = row_idx
            elif item_text == current_item and item_text != "" and current_item != "":
                # Same item in same type - continue the group
                continue
            else:
                # Different item in same type, merge previous group if needed
                if item_start != -1 and row_idx - item_start > 1 and current_item:
                    merge_cells_in_column(table, item_col_idx, item_start, row_idx - 1)

                current_item = item_text
                item_start = row_idx

        # Handle last group
        if item_start != -1 and len(table.rows) - item_start > 1 and current_item:
            merge_cells_in_column(table, item_col_idx, item_start, len(table.rows) - 1)

def fill_word_document_table(doc_path: str, output_path: str, table_data: List[List[str]],
                           table_index: Optional[int] = None,
                           target_columns: List[str] = ['类型', '检验项目', '检验方法', '质量标准'],
                           auto_merge: bool = True) -> str:
    """
    Fill Word document table with quality standards data

    Args:
        doc_path: Path to input Word document
        output_path: Path to save output document
        table_data: List of rows, each row is a list of cell values
        table_index: Specific table index to fill (None for auto-detection)
        target_columns: Expected column order
        auto_merge: Whether to automatically merge duplicate cells

    Returns:
        Success message or error description
    """
    try:
        logger.info(f"Loading document: {doc_path}")
        doc = Document(doc_path)

        # Find target table
        if table_index is None:
            table_index = find_quality_standards_table(doc)
            if table_index is None:
                return "Error: No quality standards table found in document"

        if table_index >= len(doc.tables):
            return f"Error: Table index {table_index} not found in document (only {len(doc.tables)} tables)"

        target_table = doc.tables[table_index]
        logger.info(f"Using table {table_index} with {len(target_table.rows)} rows and {len(target_table.columns)} columns")

        # Validate data
        if not table_data:
            return "Error: No table data provided"

        logger.info(f"Filling table with {len(table_data)} data rows")

        # Step 1: Clear existing table content (except header)
        clear_table_content(target_table, keep_header=True)

        # Step 2: Insert required number of rows
        num_data_rows = len(table_data)
        insert_table_rows(target_table, num_data_rows)

        # Step 3: Fill the table with data
        for i, row_data in enumerate(table_data):
            row_index = i + 1  # Skip header row
            if row_index < len(target_table.rows):
                row = target_table.rows[row_index]

                # Fill cells based on available data and columns
                max_cols = min(len(row_data), len(row.cells), len(target_columns))

                for col_idx in range(max_cols):
                    cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                    cell = row.cells[col_idx]

                    # Restore formatting for the cell content
                    restore_formatting_to_cell(cell, cell_text)

                logger.debug(f"Filled row {row_index}: {row_data[:max_cols]}")
            else:
                logger.warning(f"Skipping row {row_index}, table doesn't have enough rows")

        # Step 4: Auto-merge duplicate cells if requested
        if auto_merge:
            logger.info("Performing automatic cell merging...")
            auto_merge_duplicate_cells(target_table, target_columns[:2])  # Only merge 类型 and 检验项目

        # Step 5: Save the modified document
        doc.save(output_path)
        logger.info(f"Document saved to: {output_path}")

        return f"Successfully filled table with {len(table_data)} rows and saved to {output_path}"

    except Exception as e:
        error_msg = f"Error filling Word document table: {str(e)}"
        logger.error(error_msg)
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return error_msg

def fill_quality_standards_from_markdown(doc_path: str, output_path: str, markdown_content: str,
                                       table_index: Optional[int] = None, auto_merge: bool = True) -> str:
    """
    Fill Word document table with quality standards data from markdown content

    Args:
        doc_path: Path to input Word document
        output_path: Path to save output document
        markdown_content: Markdown table content as string
        table_index: Specific table index to fill (None for auto-detection)
        auto_merge: Whether to automatically merge duplicate cells

    Returns:
        Success message or error description
    """
    try:
        logger.info("Parsing markdown table content...")
        table_data = parse_markdown_table_from_string(markdown_content)

        if not table_data:
            return "Error: No valid table data found in markdown content"

        return fill_word_document_table(doc_path, output_path, table_data, table_index, auto_merge=auto_merge)

    except Exception as e:
        error_msg = f"Error processing markdown content: {str(e)}"
        logger.error(error_msg)
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return error_msg

def fill_quality_standards_from_file(doc_path: str, output_path: str, markdown_file_path: str,
                                   table_index: Optional[int] = None, auto_merge: bool = True) -> str:
    """
    Fill Word document table with quality standards data from markdown file

    Args:
        doc_path: Path to input Word document
        output_path: Path to save output document
        markdown_file_path: Path to markdown file containing table data
        table_index: Specific table index to fill (None for auto-detection)
        auto_merge: Whether to automatically merge duplicate cells

    Returns:
        Success message or error description
    """
    try:
        logger.info(f"Reading markdown file: {markdown_file_path}")
        table_data = parse_markdown_table_from_file(markdown_file_path)

        if not table_data:
            return "Error: No valid table data found in markdown file"

        return fill_word_document_table(doc_path, output_path, table_data, table_index, auto_merge=auto_merge)

    except Exception as e:
        error_msg = f"Error processing markdown file: {str(e)}"
        logger.error(error_msg)
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return error_msg

def fill_quality_standards_inplace(doc_path: str, markdown_content: str,
                                 table_index: Optional[int] = None, auto_merge: bool = True) -> str:
    """
    Fill Word document table with quality standards data from markdown content (modifies file in-place)

    Args:
        doc_path: Path to Word document to modify in-place
        markdown_content: Markdown table content as string
        table_index: Specific table index to fill (None for auto-detection)
        auto_merge: Whether to automatically merge duplicate cells

    Returns:
        Success message or error description
    """
    return fill_quality_standards_from_markdown(doc_path, doc_path, markdown_content, table_index, auto_merge)