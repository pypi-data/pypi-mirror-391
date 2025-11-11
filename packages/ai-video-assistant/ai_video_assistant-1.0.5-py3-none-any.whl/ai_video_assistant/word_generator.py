"""
Word document generator for the AI-Powered Video Lecture Assistant.
Creates formatted .docx files with analysis results.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from pathlib import Path
from datetime import datetime

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class WordDocumentGenerator:
    """Generates formatted Word documents from analysis results."""
    
    def __init__(self):
        """Initialize the Word document generator."""
        self.doc = None
    
    def create_document(self, result: dict, output_path: str) -> str:
        """
        Create a formatted Word document from analysis results.
        Matches the terminal output format exactly.
        
        Args:
            result: Dictionary containing analysis results
            output_path: Path to save the Word document
        
        Returns:
            Path to the created document
        """
        self.doc = Document()
        
        # Set document margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Add terminal-style header
        self._add_terminal_header()
        
        # Add summary section (terminal style)
        self._add_terminal_summary(result.get('summary', ''))
        
        # Add insights section (terminal style)
        self._add_terminal_insights(result.get('insights', []))
        
        # Add quiz section (terminal style)
        self._add_terminal_quiz(result.get('quiz', []))
        
        # Add separator line
        self._add_separator()
        
        # Save document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        
        logger.info(f"Word document created: {output_path}")
        return str(output_path)
    
    def _add_separator(self, char: str = "=", length: int = 60):
        """Add a separator line."""
        para = self.doc.add_paragraph(char * length)
        run = para.runs[0]
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(6)
    
    def _add_terminal_header(self):
        """Add terminal-style header."""
        self._add_separator()
        
        heading = self.doc.add_paragraph('VIDEO LECTURE ANALYSIS RESULTS')
        run = heading.runs[0]
        run.font.name = 'Consolas'
        run.font.size = Pt(11)
        run.font.bold = True
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.space_after = Pt(0)
        
        self._add_separator()
        self.doc.add_paragraph()  # Spacing
    
    def _add_terminal_summary(self, summary: str):
        """Add summary in terminal style."""
        # Section header
        header = self.doc.add_paragraph('📝 SUMMARY:')
        run = header.runs[0]
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.font.bold = True
        header.paragraph_format.space_after = Pt(0)
        
        # Separator
        sep = self.doc.add_paragraph('-' * 60)
        run = sep.runs[0]
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        sep.paragraph_format.space_after = Pt(6)
        
        # Content
        para = self.doc.add_paragraph(summary)
        run = para.runs[0]
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        para.paragraph_format.line_spacing = 1.2
        para.paragraph_format.space_after = Pt(12)
    
    def _add_terminal_insights(self, insights: list):
        """Add insights in terminal style."""
        # Section header
        header = self.doc.add_paragraph('💡 KEY INSIGHTS:')
        run = header.runs[0]
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.font.bold = True
        header.paragraph_format.space_after = Pt(0)
        
        # Separator
        sep = self.doc.add_paragraph('-' * 60)
        run = sep.runs[0]
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        sep.paragraph_format.space_after = Pt(6)
        
        # Insights
        for i, insight in enumerate(insights, 1):
            para = self.doc.add_paragraph(f'{i}. {insight}')
            run = para.runs[0]
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            para.paragraph_format.space_after = Pt(4)
        
        self.doc.add_paragraph()  # Spacing
    
    def _add_terminal_quiz(self, quiz: list):
        """Add quiz in terminal style."""
        # Section header
        header = self.doc.add_paragraph('❓ QUIZ QUESTIONS:')
        run = header.runs[0]
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.font.bold = True
        header.paragraph_format.space_after = Pt(0)
        
        # Separator
        sep = self.doc.add_paragraph('-' * 60)
        run = sep.runs[0]
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        sep.paragraph_format.space_after = Pt(6)
        
        # Quiz questions
        for i, q in enumerate(quiz, 1):
            # Blank line before each question (except first)
            if i > 1:
                self.doc.add_paragraph()
            
            # Question text
            q_para = self.doc.add_paragraph(f"Question {i}: {q['question']}")
            run = q_para.runs[0]
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            q_para.paragraph_format.space_after = Pt(2)
            
            # Options
            for j, option in enumerate(q['options'], 1):
                is_correct = option == q['correct_answer']
                marker = "[✓]" if is_correct else "[ ]"
                
                opt_para = self.doc.add_paragraph(f"  {marker} {j}. {option}")
                run = opt_para.runs[0]
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                
                if is_correct:
                    run.font.color.rgb = RGBColor(0, 128, 0)
                
                opt_para.paragraph_format.space_after = Pt(2)
        
        self.doc.add_paragraph()  # Final spacing


def generate_word_document(result: dict, output_path: str) -> str:
    """
    Convenience function to generate a Word document.
    
    Args:
        result: Analysis result dictionary
        output_path: Path to save the document
    
    Returns:
        Path to the created document
    """
    generator = WordDocumentGenerator()
    return generator.create_document(result, output_path)


if __name__ == "__main__":
    # Example usage
    sample_result = {
        "video_file": "sample_lecture.mp4",
        "language": "en",
        "transcription": "This is a sample transcription of the lecture content...",
        "summary": "This lecture covers the fundamentals of machine learning.",
        "insights": [
            "Machine learning enables computers to learn from data",
            "There are three main types of ML: supervised, unsupervised, and reinforcement",
            "Deep learning is a subset of machine learning"
        ],
        "quiz": [
            {
                "question": "What is machine learning?",
                "options": [
                    "A programming language",
                    "A subset of AI",
                    "A database system",
                    "An operating system"
                ],
                "correct_answer": "A subset of AI"
            }
        ]
    }
    
    generate_word_document(sample_result, "sample_output.docx")
    print("Sample document generated!")
