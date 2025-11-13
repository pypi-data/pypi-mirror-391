# -*- coding: utf-8 -*-
# ===============================================================
#
#    Copyright (C) 2024 Beike, Inc. All Rights Reserved.
#
#    @Create Author : luxu
#    @Create Time   : 2024/7/17
#    @Description   : 
#
# ===============================================================
import io
import logging

from docx import Document, ImagePart
from docx.oxml import CT_Picture
from docx.text.paragraph import Paragraph

from doc_parser.layout_parser.layout.simple_block import SimpleBlock
from services.constants import TEXT, TABLE, IMAGE
from services.layout_parse_utils import get_s3_links_for_simple_block_batch, trans_simple_block_list2string


def find_image(doc: Document, paragraph: Paragraph):
    img = paragraph._element.xpath('.//pic:pic')
    if not img:
        return
    img: CT_Picture = img[0]
    embed = img.xpath('.//a:blip/@r:embed')[0]
    related_part: ImagePart = doc.part.related_parts[embed]
    image = related_part.image
    return image


def layout_parse(file):
    simple_block_list = []

    file_stream = io.BytesIO(file)
    doc = Document(file_stream)

    for element in doc.element.body:
        try:
            if element.tag.endswith('p'):
                # 处理段落
                paragraph = doc.paragraphs[doc.element.body.index(element)]
                text = paragraph.text
                if text.strip():  # 忽略空段落
                    simple_block_list.append(SimpleBlock(type=TEXT, text=text))

                # 检查段落中的图片
                image = find_image(doc, paragraph)
                if image is not None:
                    simple_block_list.append(SimpleBlock(type=IMAGE, image_bytes=image.blob))
            elif element.tag.endswith('tbl'):
                # 处理表格
                table = element
                table_text = ""
                for row in table.findall('.//w:tr', doc.element.nsmap):
                    for cell in row.findall('.//w:tc', doc.element.nsmap):
                        table_text = " | ".join([table_text, next(cell.itertext(), '')])
                if table_text:
                    simple_block_list.append(SimpleBlock(type=TABLE, text=table_text))
        except Exception as e:
            logging.error(f"处理元素时出错，type: {element}，errmsg：{str(e)}")
            continue

    # SimpleBlock的list批量获取S3链接，并返回目标结构
    result_json = get_s3_links_for_simple_block_batch(simple_block_list)
    result_text = trans_simple_block_list2string(result_json)
    return result_json, result_text


def get_paragraph_count(file):
    file_stream = io.BytesIO(file)
    doc = Document(file_stream)
    paragraph_count = len(doc.paragraphs)
    return paragraph_count
