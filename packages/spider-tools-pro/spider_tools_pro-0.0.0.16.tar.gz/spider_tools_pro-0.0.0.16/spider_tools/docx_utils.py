import os
import re
from io import BytesIO
from typing import Dict, List, Tuple

from loguru import logger
from docx import Document
import tempfile
from pathlib import Path
import pythoncom
import win32com.client
import requests

def to_docx(doc_bytes: bytes) -> bytes:
    """
    输入: DOC 的二进制内容（bytes）
    输出: 转换后的 DOCX 二进制内容（bytes）

    - Windows + Office 可用时：使用 Word COM 高保真转换
    - 否则：自动降级为纯文本方案（仅保留文本）
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = Path(tmpdir) / "input.doc"
        out_docx_path = Path(tmpdir) / "output.docx"
        doc_path.write_bytes(doc_bytes)

        # 1) 尝试 Word COM（Windows + 安装 Office）
        try:
            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(doc_path))
            doc.SaveAs(str(out_docx_path), FileFormat=16)  # 16 = docx
            doc.Close(False)
            word.Quit()
            pythoncom.CoUninitialize()
            return out_docx_path.read_bytes()
        except Exception as e:
            print(f"Word COM 转换失败: {e}")

def remove_blank(text: str) -> str:
    """移除空格与零宽空格。"""
    return re.sub(r"[ \u200B]+", "", text)


# def extract_docx(file_path: str) -> List[Dict]:
#     """
#     解析 .docx 文件为片段列表：段落与表格。
#
#     返回的每个片段包含：
#     - heading_level: int|None 标题级别（Heading 1 → 1），非标题为 None
#     - font_size: float|None 段落字号（pt）
#     - content: str 文本内容（表格按 Markdown 行输出）
#     表格以 table_start → (table 行) → table_end 作为范围标记。
#     """
#     doc = Document(file_path)
#     fragments: List[Dict] = []
#
#     for block in doc.element.body.iterchildren():
#         if block.tag.endswith("p"):
#             para = next((p for p in doc.paragraphs if p._element == block), None)
#             if not para:
#                 continue
#
#             style_name = para.style.name if para.style else ""
#             heading_level = None
#             if style_name and style_name.startswith("Heading"):
#                 try:
#                     heading_level = int(style_name.replace("Heading", "").strip())
#                 except Exception:
#                     heading_level = None
#
#             texts: List[str] = []
#             font_size = None
#             for run in para.runs:
#                 txt = run.text.strip()
#                 if run.font and run.font.size:
#                     try:
#                         font_size = run.font.size.pt
#                     except Exception:
#                         font_size = None
#                 if txt:
#                     texts.append(txt)
#
#             if texts:
#                 fragments.append({
#                     "heading_level": heading_level,
#                     "font_size": font_size,
#                     "content": remove_blank(" ".join(texts)),
#                 })
#
#         elif block.tag.endswith("tbl"):
#             tbl = next((t for t in doc.tables if t._element == block), None)
#             if not tbl:
#                 continue
#
#             fragments.append({"heading_level": "table_start", "font_size": None, "content": "\n"})
#
#             for r_idx, row in enumerate(tbl.rows):
#                 row_texts: List[str] = []
#                 for cell in row.cells:
#                     cell_lines: List[str] = []
#                     for para in cell.paragraphs:
#                         for run in para.runs:
#                             if run.text.strip():
#                                 cell_lines.append(run.text.strip())
#                     row_texts.append(re.sub(r"\|+", " ", " <br> ".join(cell_lines)))
#
#                 md_row = f"|{remove_blank('|'.join(row_texts))}|"
#                 fragments.append({"heading_level": "table", "font_size": None, "content": md_row})
#
#                 if r_idx == 0:
#                     cols = len(row_texts)
#                     fragments.append({
#                         "heading_level": "table",
#                         "font_size": None,
#                         "content": f"|{remove_blank('|'.join(['--------'] * cols))}|",
#                     })
#
#             fragments.append({"heading_level": "table_end", "font_size": None, "content": "\n"})
#
#     return fragments
def extract_docx(file_path: str = None, content: bytes = None) -> List[Dict]:
    """
    解析 .docx 文件（路径或二进制内容）为片段列表：段落与表格。

    参数：
    - file_path: str 可选，docx 文件路径
    - content: bytes 可选，docx 二进制内容（如内存中的文件内容）
    注：file_path 和 content 必须传入其中一个

    返回的每个片段包含：
    - heading_level: int|None 标题级别（Heading 1 → 1），非标题为 None
    - font_size: float|None 段落字号（pt）
    - content: str 文本内容（表格按 Markdown 行输出）
    表格以 table_start → (table 行) → table_end 作为范围标记。
    """
    # 新增：校验输入，确保必传其一
    if file_path is None and content is None:
        raise ValueError("必须传入 file_path（文件路径）或 content（二进制内容）")
    if file_path is not None and content is not None:
        raise ValueError("file_path 和 content 不能同时传入，二选一即可")
    if file_path is not None:
        doc = Document(file_path)
    else:
        doc = Document(BytesIO(content))
    fragments: List[Dict] = []
    for block in doc.element.body.iterchildren():
        if block.tag.endswith("p"):
            para = next((p for p in doc.paragraphs if p._element == block), None)
            if not para:
                continue

            style_name = para.style.name if para.style else ""
            heading_level = None
            if style_name and style_name.startswith("Heading"):
                try:
                    heading_level = int(style_name.replace("Heading", "").strip())
                except Exception:
                    heading_level = None

            texts: List[str] = []
            font_size = None
            for run in para.runs:
                txt = run.text.strip()
                if run.font and run.font.size:
                    try:
                        font_size = run.font.size.pt
                    except Exception:
                        font_size = None
                if txt:
                    texts.append(txt)

            if texts:
                fragments.append({
                    "heading_level": heading_level,
                    "font_size": font_size,
                    "content": remove_blank(" ".join(texts)),
                })

        elif block.tag.endswith("tbl"):
            tbl = next((t for t in doc.tables if t._element == block), None)
            if not tbl:
                continue

            fragments.append({"heading_level": "table_start", "font_size": None, "content": "\n"})

            for r_idx, row in enumerate(tbl.rows):
                row_texts: List[str] = []
                for cell in row.cells:
                    cell_lines: List[str] = []
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text.strip():
                                cell_lines.append(run.text.strip())
                    row_texts.append(re.sub(r"\|+", " ", " <br> ".join(cell_lines)))

                md_row = f"|{remove_blank('|'.join(row_texts))}|"
                fragments.append({"heading_level": "table", "font_size": None, "content": md_row})

                if r_idx == 0:
                    cols = len(row_texts)
                    fragments.append({
                        "heading_level": "table",
                        "font_size": None,
                        "content": f"|{remove_blank('|'.join(['--------'] * cols))}|",
                    })

            fragments.append({"heading_level": "table_end", "font_size": None, "content": "\n"})

    return fragments

def judge_head_level(text: str) -> bool:
    """判断是否为中文编号的一节标题，如 '一、' '二.'。"""
    chinese_num_pattern = re.compile(r"[一二三四五六七八九十]+[、.]")
    return bool(chinese_num_pattern.match(text))


def deal_text(text_dict_list: List[Dict]) -> Dict:
    """
    将 extract_docx 片段聚合为结构化结果：
    返回：
    {
      'title': [大标题...],
      'content': { '一级标题': [内容行...], ... }
    }
    """
    import copy

    section_list: List[Dict] = []
    section_dict: Dict = {}
    current_level = '前言'
    last_font_size = 0
    for text_dict in text_dict_list:
        font_size = text_dict.get('font_size')
        content = text_dict.get('content')
        if font_size == 22.0:
            if last_font_size == font_size:
                section_dict.setdefault('title', []).append(content)
            else:
                if section_dict:
                    if 'title' in section_dict:
                        section_list.append(copy.deepcopy(section_dict))
                    section_dict = {}
                    current_level = '前言'
                section_dict.setdefault('title', []).append(content)
        else:
            section_dict.setdefault('content', {})
            if content and judge_head_level(content):
                current_level = content
                section_dict['content'][content] = []
            else:
                if current_level not in section_dict['content']:
                    section_dict['content'][current_level] = []
                section_dict['content'][current_level].append(content)
        last_font_size = font_size

    section_list.append(copy.deepcopy(section_dict))
    # 合并标题段落为一级键
    section = copy.deepcopy(section_list[0]) if section_list else {}
    if len(section_list):
        for temp in section_list[1:-1]:
            content_key = temp.get('title')
            if content_key:
                content_value: List[str] = []
                for key, value in temp.get('content', {}).items():
                    content_value.append(key)
                    content_value.extend(value)
                section.setdefault('content', {})[''.join(content_key)] = content_value
    return section


def convert_doc_to_docx_pure(doc_path: str, out_docx_path: str) -> str:
    """
    将 .doc 粗略转换为 .docx（不依赖外部软件）：
    - 方法：读取二进制内容，按 'latin-1' / 'gbk' 兼容方式尽力提取文本，
      然后用 python-docx 生成简易 docx（仅保留纯文本）。
    - 注意：这是降级方案，格式/表格/图片将丢失。
    返回生成的 docx 路径。
    """
    base_dir = os.path.dirname(out_docx_path)
    if base_dir and not os.path.exists(base_dir):
        os.makedirs(base_dir, exist_ok=True)

    with open(doc_path, 'rb') as f:
        raw = f.read()
    # 尝试多种解码获取最大化文本
    text = None
    for enc in ('utf-8', 'gbk', 'latin-1'):  # 宽松策略
        try:
            text = raw.decode(enc, errors='ignore')
            if text and len(text.strip()) > 0:
                break
        except Exception:
            continue
    if text is None:
        text = ''

    docx = Document()
    for line in re.split(r"[\r\n]+", text):
        docx.add_paragraph(line)
    docx.save(out_docx_path)
    logger.warning(".doc 转换采用纯文本降级方案，格式已丢失: {} -> {}", doc_path, out_docx_path)
    return out_docx_path


def collect_docx_files_pure(src_dir: str, dest_dir: str) -> None:
    """
    批量收集 .doc/.docx，且在无外部软件的前提下处理：
    - .docx: 直接复制到目标目录
    - .doc: 采用 convert_doc_to_docx_pure 生成简化版 .docx 到目标目录
    """
    if not os.path.exists(dest_dir):
        os.makedirs(dest_dir, exist_ok=True)

    for root, _, files in os.walk(src_dir):
        for file in files:
            src = os.path.join(root, file)
            if file.lower().endswith('.docx'):
                target = os.path.join(dest_dir, os.path.basename(src))
                try:
                    from shutil import copy2
                    copy2(src, target)
                except Exception as e:
                    logger.error(f"复制失败: {src} -> {target}, {e}")
            elif file.lower().endswith('.doc') and not file.lower().endswith('.docx'):
                new_name = os.path.splitext(os.path.basename(file))[0] + '.docx'
                target = os.path.join(dest_dir, new_name)
                try:
                    convert_doc_to_docx_pure(src, target)
                except Exception as e:
                    logger.error(f".doc 转换失败: {src} -> {target}, {e}")


