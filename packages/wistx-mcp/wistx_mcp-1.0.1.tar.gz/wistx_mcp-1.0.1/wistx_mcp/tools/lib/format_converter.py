"""Format converter for report generation (Markdown, HTML, PDF, DOCX)."""

import logging
from io import BytesIO
from pathlib import Path
from typing import Any

import markdown
from docx import Document
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)

try:
    from markdownify import markdownify as md
    MARKDOWNIFY_AVAILABLE = True
except ImportError:
    MARKDOWNIFY_AVAILABLE = False
    logger.warning("markdownify not available. HTML to Markdown conversion will be limited.")

WEASYPRINT_AVAILABLE = False
WEASYPRINT_HTML = None
WEASYPRINT_CSS = None

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
    WEASYPRINT_HTML = HTML
    WEASYPRINT_CSS = CSS
except (ImportError, OSError) as e:
    WEASYPRINT_AVAILABLE = False
    logger.warning("WeasyPrint not available. PDF generation will be disabled. Error: %s", e)


class FormatConverter:
    """Converts content between formats (Markdown, HTML, PDF, DOCX)."""

    def __init__(self):
        """Initialize format converter."""
        self.md_extensions = [
            "extra",
            "codehilite",
            "tables",
            "fenced_code",
            "toc",
        ]

    def markdown_to_html(
        self,
        markdown_content: str,
        styles: dict[str, Any] | None = None,
    ) -> str:
        """Convert Markdown to HTML.

        Args:
            markdown_content: Markdown content
            styles: CSS styles dictionary

        Returns:
            HTML content
        """
        html_content = markdown.markdown(
            markdown_content,
            extensions=self.md_extensions,
        )

        css_styles = self._build_css(styles or {})

        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        {css_styles}
    </style>
</head>
<body>
    {html_content}
</body>
</html>"""

        return html

    def markdown_to_pdf(
        self,
        markdown_content: str,
        styles: dict[str, Any] | None = None,
        branding: dict[str, Any] | None = None,
    ) -> bytes:
        """Convert Markdown to PDF.

        Args:
            markdown_content: Markdown content
            styles: CSS styles dictionary
            branding: Branding configuration

        Returns:
            PDF bytes

        Raises:
            ValueError: If WeasyPrint not available
        """
        if not WEASYPRINT_AVAILABLE or WEASYPRINT_HTML is None or WEASYPRINT_CSS is None:
            raise ValueError(
                "WeasyPrint not available. Install with: pip install weasyprint "
                "and system dependencies (see https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#installation)"
            )

        html_content = self.markdown_to_html(markdown_content, styles)

        css_styles = self._build_css(styles or {}, branding)

        try:
            pdf_bytes = WEASYPRINT_HTML(string=html_content).write_pdf(
                stylesheets=[WEASYPRINT_CSS(string=css_styles)]
            )
            return pdf_bytes
        except (RuntimeError, OSError, ValueError) as e:
            logger.error("PDF generation failed: %s", e)
            raise ValueError(f"PDF generation failed: {e}") from e
        except Exception as e:
            logger.error("Unexpected error during PDF generation: %s", e)
            raise ValueError(f"PDF generation failed: Unexpected error") from e

    def markdown_to_docx(
        self,
        markdown_content: str,
        branding: dict[str, Any] | None = None,
    ) -> bytes:
        """Convert Markdown to DOCX.

        Args:
            markdown_content: Markdown content
            branding: Branding configuration

        Returns:
            DOCX bytes
        """
        doc = Document()

        if branding:
            if branding.get("logo_path"):
                try:
                    logo_path = Path(branding["logo_path"])
                    if logo_path.exists():
                        doc.add_picture(str(logo_path), width=Inches(2))
                except Exception as e:
                    logger.warning("Failed to add logo: %s", e)

            if branding.get("company_name"):
                title = doc.add_heading(branding["company_name"], 0)
                title.alignment = 1

        lines = markdown_content.split("\n")
        current_paragraph = None

        for line in lines:
            line = line.strip()

            if not line:
                if current_paragraph:
                    current_paragraph = None
                continue

            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("#### "):
                doc.add_heading(line[5:], level=4)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("```"):
                continue
            else:
                current_paragraph = doc.add_paragraph(line)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def html_to_markdown(self, html_content: str) -> str:
        """Convert HTML to Markdown.

        Args:
            html_content: HTML content

        Returns:
            Markdown content
        """
        if not MARKDOWNIFY_AVAILABLE:
            logger.warning("markdownify not available. Returning HTML content.")
            return html_content
        return md(html_content)

    def _build_css(
        self,
        styles: dict[str, Any],
        branding: dict[str, Any] | None = None,
    ) -> str:
        """Build CSS from styles dictionary.

        Args:
            styles: Styles dictionary
            branding: Branding configuration

        Returns:
            CSS string
        """
        css_parts = []

        primary_color = branding.get("primary_color", "#2563eb") if branding else "#2563eb"
        secondary_color = branding.get("secondary_color", "#64748b") if branding else "#64748b"
        font_family = branding.get("font_family", "Arial, sans-serif") if branding else "Arial, sans-serif"

        css_parts.append(f"""
            body {{
                font-family: {font_family};
                line-height: 1.6;
                color: #333;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
            }}
            h1 {{
                color: {primary_color};
                border-bottom: 2px solid {primary_color};
                padding-bottom: 10px;
            }}
            h2 {{
                color: {secondary_color};
                margin-top: 30px;
            }}
            code {{
                background-color: #f4f4f4;
                padding: 2px 6px;
                border-radius: 3px;
            }}
            pre {{
                background-color: #f4f4f4;
                padding: 15px;
                border-radius: 5px;
                overflow-x: auto;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin: 20px 0;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 12px;
                text-align: left;
            }}
            th {{
                background-color: {primary_color};
                color: white;
            }}
        """)

        if styles:
            for selector, properties in styles.items():
                props_str = "; ".join(f"{k}: {v}" for k, v in properties.items())
                css_parts.append(f"{selector} {{ {props_str} }}")

        return "\n".join(css_parts)

