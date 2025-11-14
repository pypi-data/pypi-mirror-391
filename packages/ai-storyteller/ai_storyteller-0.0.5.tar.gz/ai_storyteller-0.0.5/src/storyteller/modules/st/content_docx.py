from __future__ import annotations

import io
import re
import datetime

import requests

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from storyteller.modules.st.content import break_text_into_paragraphs

__all__ = (
    "create_docx",
)


def _fetch_image_bytes(url_or_path: str) -> io.BytesIO | None:
    """Fetch image from URL or local path and return as BytesIO, or None on failure."""
    try:
        if url_or_path.startswith("http://") or url_or_path.startswith("https://"):
            resp = requests.get(url_or_path, timeout=20)
            resp.raise_for_status()
            return io.BytesIO(resp.content)
        # Local path fallback
        with open(url_or_path, "rb") as fh:
            return io.BytesIO(fh.read())
    except Exception:
        return None

def _usable_page_width_inches(document: Document) -> float:
    """Compute usable page width in Inches based on the first section."""
    section = document.sections[0]
    page_width_inches = section.page_width.inches
    left_margin_inches = section.left_margin.inches
    right_margin_inches = section.right_margin.inches
    return max(0.0, page_width_inches - left_margin_inches - right_margin_inches)

def _add_centered_picture(document: Document, img_stream: io.BytesIO, width_inches: float) -> None:
    """Insert centered image with given width in inches."""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_stream, width=Inches(width_inches))

def _add_centered_text(document: Document, text: str, bold: bool = False, font_size_pt: int | None = None) -> None:
    """Add a centered paragraph with optional bold and font size."""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bool(bold)
    if font_size_pt:
        run.font.size = Pt(font_size_pt)

def _add_justified_paragraph(document: Document, text: str, space_before_pt: int = 0, space_after_pt: int = 12, line_spacing: float = 1.5) -> None:
    """Add a justified paragraph with spacing and line-height."""
    p = document.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    if space_before_pt:
        pf.space_before = Pt(space_before_pt)
    if space_after_pt:
        pf.space_after = Pt(space_after_pt)
    # Set line spacing
    pf.line_spacing = line_spacing

def create_docx(
    title: str,
    text: str,
    images: list[str],
    cover_image_url: str,
    back_cover_image_url: str,
    output_path: str,
    page_margins_inches: float = 1.0,
    text_block_width_ratio: float = 0.75,
    image_width_ratio: float = 0.95,
    back_cover_small_ratio: float = 0.60,
) -> str:
    """
    Create a DOCX that mirrors the markdown_with_images logic:
    - Cover: H1 title + cover image (full width).
    - Interleave paragraphs with images; leftover images appended after all paragraphs.
    - 'The End' page with small back-cover image.
    - Final page with full back-cover image.

    Returns the output_path on success.
    """
    # Create document and set basic page geometry
    doc = Document()
    section = doc.sections[0]
    margin = Inches(page_margins_inches)
    section.left_margin = margin
    section.right_margin = margin
    section.top_margin = margin
    section.bottom_margin = margin

    usable_width = _usable_page_width_inches(doc)
    # Text block width inside usable page width
    text_block_width = max(0.0, usable_width * text_block_width_ratio)
    # Image widths
    regular_image_width = max(0.0, usable_width * image_width_ratio)
    cover_full_width = usable_width
    back_cover_small_width = max(0.0, usable_width * back_cover_small_ratio)
    back_cover_full_width = usable_width

    # Title style (centered, bold, larger)
    # Note: python-docx doesn't have built-in 'Heading 1' centering by default,
    # so we add an explicit centered paragraph.
    _add_centered_text(doc, title, bold=True, font_size_pt=24)

    # Cover image
    cover_stream = _fetch_image_bytes(cover_image_url)
    if cover_stream:
        _add_centered_picture(doc, cover_stream, cover_full_width)

    # New page
    doc.add_page_break()

    # --- Attribution page as page 2 (unchanged) ---
    current_year = datetime.datetime.now().year  # Compute year once here
    doc.add_paragraph()
    _add_centered_text(doc, "Created by:", bold=False, font_size_pt=12)
    _add_centered_text(doc, "Artur Barseghyan", bold=False, font_size_pt=12)
    _add_centered_text(doc, "Dale Richardson", bold=False, font_size_pt=12)
    doc.add_paragraph()  # For spacing
    _add_centered_text(doc, "Publisher: AI Storyteller", bold=False, font_size_pt=12)
    _add_centered_text(doc, f"Year: {current_year}", bold=False, font_size_pt=12)
    doc.add_page_break()
    # --- End attribution page ---

    # Prepare paragraphs list akin to markdown_with_images
    normalized_text = break_text_into_paragraphs(text, nb_paragraphs=len(images))
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", normalized_text.strip()) if p.strip()]

    # Interleave paragraphs and images
    img_idx = 0
    for i, para in enumerate(paragraphs):
        # Text block: we do not actually constrain to a fixed width,
        # but we simulate layout by spacing and justification.
        _add_justified_paragraph(doc, para, space_before_pt=0, space_after_pt=6, line_spacing=1.5)

        # Insert image after the paragraph if available
        if img_idx < len(images):
            stream = _fetch_image_bytes(images[img_idx])
            if stream:
                _add_centered_picture(doc, stream, regular_image_width)
                doc.add_page_break()  # NEW: break after every *story* image
            img_idx += 1
        else:
            # Add some vertical whitespace between blocks only if no image inserted
            doc.add_paragraph()  # Blank line

        # Optionally add a page break per block if desired; we stay on the same page
        # to let Word paginate naturally, as markdown_with_images does not force page per block.

    # Append leftover images (if any)
    while img_idx < len(images):
        stream = _fetch_image_bytes(images[img_idx])
        if stream:
            _add_centered_picture(doc, stream, regular_image_width)
            doc.add_page_break()  # NEW: break after each leftover *story* image
        img_idx += 1
        # No extra blank paragraph here; the page break is sufficient

    # "The End" page
    doc.add_page_break()
    _add_centered_text(doc, "The End", bold=True, font_size_pt=20)

    # Small back-cover image on the 'The End' page
    bc_small_stream = _fetch_image_bytes(back_cover_image_url)
    if bc_small_stream:
        _add_centered_picture(doc, bc_small_stream, back_cover_small_width)

    # Final back-cover page (full-width image)
    doc.add_page_break()
    bc_full_stream = _fetch_image_bytes(back_cover_image_url)
    if bc_full_stream:
        _add_centered_picture(doc, bc_full_stream, back_cover_full_width)

    # Footer note on the last page (centered, small)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI Storyteller")
    run.font.size = Pt(10)

    # Save
    doc.save(output_path)
    return output_path
